#!/usr/bin/env python3
"""Post-process DOCX after PDF conversion.

1. Unwrap bullet/number lists that pdf2docx wrongly emitted as tables.
2. Repair fused words at paragraph level (pdf2docx often omits spaces between
   glyph runs).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import wordninja
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BULLET_PREFIX = re.compile(
    r"^[\s]*(?:[•●○◦▪▫■□◆◇►▻➢·]|[-–—*+]\s|\d+[.)]\s|[A-Za-z][.)]\s)"
)

# Longer product/stack tokens protected before CamelCase / digit splitting.
GLOSSARY = (
    "PostgreSQL",
    "GuidelyEdu",
    "Laravel",
    "Capacitor",
    "TypeScript",
    "JavaScript",
    "MongoDB",
    "MySQL",
    "SQLite",
    "Redis",
    "Docker",
    "Kubernetes",
    "GitHub",
    "GitLab",
    "Bitbucket",
    "OpenAPI",
    "GraphQL",
    "Protégé",
    "Protege",
    "Toolchain",
    "Vue",
)

SHORT_GLOSSARY = (
    "SRE",
    "OCR",
    "PDF",
    "API",
    "CLI",
    "SDK",
    "JWT",
    "PRD",
)

URL_PATTERN = re.compile(r"https?://\S+|www\.\S+")
EMAIL_PATTERN = re.compile(r"\b[\w.+-]+@[\w.-]+\.\w+\b")
EXTENSION_PATTERN = re.compile(
    r"\.(?:md|txt|pdf|docx?|xlsx?|pptx?|csv|json|ya?ml|xml|html?|css|js|ts|tsx|jsx|py|php|"
    r"java|go|rs|sql|env|log|png|jpe?g|gif|svg|webp|zip|tar|gz|tgz|rar|7z|sh|bat|ps1)\b",
    re.IGNORECASE,
)
LONG_ALPHA = re.compile(r"[A-Za-z]{6,}")
PLACEHOLDER = re.compile(r"¤\d+¤")


def _protect(text: str) -> tuple[str, dict[str, str]]:
    protected: dict[str, str] = {}

    def stash(value: str) -> str:
        key = f"¤{len(protected)}¤"
        protected[key] = value
        return key

    def stash_match(match: re.Match[str]) -> str:
        return stash(match.group(0))

    text = URL_PATTERN.sub(stash_match, text)
    text = EMAIL_PATTERN.sub(stash_match, text)
    text = EXTENSION_PATTERN.sub(stash_match, text)

    for term in sorted(GLOSSARY, key=len, reverse=True):
        if term in text:
            text = text.replace(term, stash(term))

    for term in SHORT_GLOSSARY:
        text = re.sub(rf"\b{re.escape(term)}\b", lambda match: stash(match.group(0)), text)

    return text, protected


def _restore(text: str, protected: dict[str, str]) -> str:
    for index in range(len(protected) - 1, -1, -1):
        key = f"¤{index}¤"
        text = text.replace(key, protected[key])
    return text


def _merge_crumbs(parts: list[str]) -> list[str]:
    cleaned: list[str] = []
    for part in parts:
        if (
            cleaned
            and len(part) == 1
            and part not in {"a", "i"}
            and len(cleaned[-1]) <= 3
        ):
            cleaned[-1] += part
        else:
            cleaned.append(part)
    return cleaned


def _split_long_token(token: str) -> str:
    if PLACEHOLDER.fullmatch(token) or not LONG_ALPHA.fullmatch(token):
        return token

    parts = _merge_crumbs(wordninja.split(token.lower()))
    if len(parts) < 2:
        return token

    if any(len(part) == 1 and part not in {"a", "i"} for part in parts):
        return token

    if token.isupper():
        return " ".join(part.upper() for part in parts)
    if token[0].isupper():
        rebuilt = " ".join(parts)
        return rebuilt[:1].upper() + rebuilt[1:]
    return " ".join(parts)


def respace_paragraph_text(text: str) -> str:
    if not text or text.isspace():
        return text

    text, protected = _protect(text)

    # Keep glossary/file placeholders from gluing to neighbouring letters/digits.
    text = re.sub(r"(¤\d+¤)([A-Za-z0-9])", r"\1 \2", text)
    text = re.sub(r"([A-Za-z0-9])(¤\d+¤)", r"\1 \2", text)

    # Punctuation glued to words.
    text = re.sub(r"([,.;:!?%/])([A-Za-z])", r"\1 \2", text)
    text = re.sub(r"([A-Za-z0-9])([(])", r"\1 \2", text)
    text = re.sub(r"([)])([A-Za-z0-9])", r"\1 \2", text)
    text = re.sub(r"([;:])([A-Za-z])", r"\1 \2", text)

    # Possessives / curly apostrophes stuck between words.
    text = re.sub(r"([A-Za-z])['’]s([A-Za-z])", r"\1's \2", text)

    # CamelCase and letter/digit boundaries.
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    text = re.sub(r"([A-Za-z])(\d)", r"\1 \2", text)
    text = re.sub(r"(\d)([A-Za-z])", r"\1 \2", text)

    text = LONG_ALPHA.sub(lambda match: _split_long_token(match.group(0)), text)

    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r" +([,.;:!?])", r"\1", text)

    # Don't leave a gap before file-extension placeholders (".md", ".pdf", ...).
    for key, value in protected.items():
        if value.startswith("."):
            text = text.replace(f" {key}", key)

    return _restore(text, protected)



def _cell_text(cell) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _normalize_list_item(text: str) -> str:
    text = text.strip()
    text = BULLET_PREFIX.sub("", text).strip()
    return text


def _row_list_text(row) -> str | None:
    cells = row.cells
    if not cells:
        return None

    if len(cells) == 1:
        return _cell_text(cells[0])

    if len(cells) == 2:
        left = _cell_text(cells[0])
        right = _cell_text(cells[1])
        if not left and not right:
            return ""
        if not right:
            return left
        if not left or len(left) <= 3 or BULLET_PREFIX.match(left + " "):
            if left and BULLET_PREFIX.match(left + " "):
                return f"{left} {right}".strip()
            if left in {"•", "●", "○", "◦", "▪", "▫", "-", "–", "—", "*", "·"}:
                return f"{left} {right}".strip()
            return right
        return None

    return None


def _table_looks_like_list(table) -> bool:
    rows = table.rows
    if len(rows) < 2:
        return False
    if len(rows[0].cells) > 2:
        return False

    texts: list[str] = []
    for row in rows:
        item = _row_list_text(row)
        if item is None:
            return False
        if item:
            texts.append(item)

    if len(texts) < 2:
        return False

    bulletish = sum(1 for text in texts if BULLET_PREFIX.match(text))
    return (bulletish / len(texts)) >= 0.45


def _list_paragraph_element(text: str):
    paragraph = OxmlElement("w:p")
    props = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "ListBullet")
    props.append(style)
    paragraph.append(props)

    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    paragraph.append(run)
    return paragraph


def unwrap_false_list_tables(document: Document) -> int:
    """Replace bullet-like 1–2 column tables with real list paragraphs."""
    changed = 0
    # Snapshot because we mutate the body while walking.
    for table in list(document.tables):
        if not _table_looks_like_list(table):
            continue

        tbl = table._tbl
        parent = tbl.getparent()
        if parent is None:
            continue

        index = list(parent).index(tbl)
        items = []
        for row in table.rows:
            item = _row_list_text(row)
            if item:
                items.append(_normalize_list_item(item) or item)

        for offset, item in enumerate(items):
            parent.insert(index + offset, _list_paragraph_element(item))

        parent.remove(tbl)
        changed += 1

    return changed


def _fix_paragraph(paragraph) -> bool:
    runs = paragraph.runs
    if not runs:
        return False

    original = "".join(run.text for run in runs)
    fixed = respace_paragraph_text(original)
    if fixed == original:
        return False

    runs[0].text = fixed
    for run in runs[1:]:
        run.text = ""
    return True


def fix_docx(input_path: Path, output_path: Path) -> int:
    document = Document(str(input_path))
    changed = unwrap_false_list_tables(document)

    for paragraph in document.paragraphs:
        if _fix_paragraph(paragraph):
            changed += 1

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _fix_paragraph(paragraph):
                        changed += 1

    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if _fix_paragraph(paragraph):
                    changed += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return changed


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Input DOCX path")
    parser.add_argument("-o", "--output", type=Path, required=True, help="Output DOCX path")
    args = parser.parse_args(argv)

    if not args.input.is_file():
        print(f"Input not found: {args.input}", file=sys.stderr)
        return 1

    changed = fix_docx(args.input, args.output)
    print(f"Repaired {changed} paragraph/table issue(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
